from docx import Document
from docx.enum.text import WD_BREAK
import re

def replace_text_in_docx(doc: Document, data: dict):
  for paragraph in doc.paragraphs:
    for key, value in data.items():
      for run in paragraph.runs:
        run.text = run.text.replace(f"<{key}>", str(value))
  for table in doc.tables:
    for row in table.rows:
      for cell in row.cells:
        for key, value in data.items():
          cell.text = cell.text.replace(f"<{key}>", str(value))

def replace_text_in_docx_robust(doc: Document, data: dict):
    """
    Reemplaza placeholders en un documento DOCX, manejando casos donde
    los placeholders están divididos en múltiples 'runs'.
    """
    # Expresión regular para encontrar todos los placeholders del tipo <key>
    regex = re.compile(r"<(" + "|".join(data.keys()) + r")>")

    # --- Reemplazo en Párrafos ---
    for p in doc.paragraphs:
        # Si no hay un '<' en el texto del párrafo, es muy improbable que haya un placeholder.
        # Esto es una optimización para evitar procesar párrafos innecesariamente.
        if '<' not in p.text:
            continue

        # Combina el texto de todos los 'runs' en el párrafo
        full_text = "".join(run.text for run in p.runs)
        
        # Si se encuentra un placeholder en el texto combinado
        if regex.search(full_text):
            # Reemplaza todos los placeholders encontrados en el texto combinado
            # Usamos una función lambda para buscar el valor correspondiente en el diccionario data
            new_text = regex.sub(lambda m: str(data.get(m.group(1), m.group(0))), full_text)
            
            # Borra el contenido de todos los 'runs' existentes en el párrafo
            for run in p.runs:
                run.text = ''
            
            # Añade un nuevo 'run' al principio con todo el texto modificado,
            # manteniendo el formato del primer 'run' original.
            p.add_run(new_text)

    # --- Reemplazo en Tablas ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # La lógica para las celdas es la misma que para los párrafos.
                # Una celda puede contener múltiples párrafos.
                for p in cell.paragraphs:
                    if '<' not in p.text:
                        continue
                    
                    full_text = "".join(run.text for run in p.runs)
                    if regex.search(full_text):
                        new_text = regex.sub(lambda m: str(data.get(m.group(1), m.group(0))), full_text)
                        for run in p.runs:
                            run.text = ''
                        p.add_run(new_text)

# def replace_text_in_docx(doc: Document, data: dict):
#     # Reemplazo en párrafos
#     for paragraph in doc.paragraphs:
#         runs = paragraph.runs
#         full_text = ''.join(run.text for run in runs)
#         for key, value in data.items():
#             marker = f"<{key}>"
#             if marker in full_text:
#                 # Si el marcador está partido, reconstruir los runs
#                 new_text = full_text.replace(marker, str(value))
#                 # Borrar todos los runs
#                 for run in runs:
#                     run.text = ''
#                 # Crear un solo run con el texto reemplazado (con formato del primer run)
#                 if runs:
#                     runs[0].text = new_text
#                 break  # Solo reemplazar una vez por párrafo

#         # Si no está partido, reemplazar en cada run (conserva formato)
#         for key, value in data.items():
#             marker = f"<{key}>"
#             for run in runs:
#                 if marker in run.text:
#                     run.text = run.text.replace(marker, str(value))

#     # Reemplazo en tablas (igual que antes)
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     runs = paragraph.runs
#                     full_text = ''.join(run.text for run in runs)
#                     for key, value in data.items():
#                         marker = f"<{key}>"
#                         if marker in full_text:
#                             new_text = full_text.replace(marker, str(value))
#                             for run in runs:
#                                 run.text = ''
#                             if runs:
#                                 runs[0].text = new_text
#                             break
#                     for key, value in data.items():
#                         marker = f"<{key}>"
#                         for run in runs:
#                             if marker in run.text:
#                                 run.text = run.text.replace(marker, str(value))

def insert_page_break_after_paragraph(doc: Document, search_text):
  for paragraph in doc.paragraphs:
    if search_text in paragraph.text:
      run = paragraph.add_run()
      run.add_break(WD_BREAK.PAGE)

def insert_page_break_before_paragraph(doc: Document, search_text: str):
  count = 0
  for i, paragraph in enumerate(doc.paragraphs):
    if search_text in paragraph.text:
      count += 1
      if count == 1:
        continue
      prior_paragraph = doc.paragraphs[i - 1] if i > 0 else None
      if prior_paragraph:
        run = prior_paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)